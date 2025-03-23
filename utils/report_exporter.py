# -*- coding: utf-8 -*-
import os
import io
import zipfile
import tempfile
import logging
import traceback
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple, Union

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ReportExporter:
    """
    报告导出工具类，负责处理Word文档模板和生成学生报告
    """
    
    def __init__(self, templates_dir: str = "templates/docx"):
        """
        初始化报告导出器
        
        Args:
            templates_dir: 存放模板文件的目录路径
        """
        self.templates_dir = templates_dir
        # 确保模板目录存在
        os.makedirs(os.path.join(templates_dir, "custom"), exist_ok=True)
        
        # 内置默认模板数据（作为最后备选）
        self.has_default_backup = False
        try:
            # 导入docx库创建一个基础模板
            import docx
            self.default_template = self._create_default_template()
            self.has_default_backup = True
            logger.info("成功创建内置默认模板")
        except Exception as e:
            logger.warning(f"无法创建内置默认模板: {str(e)}")
    
    def _create_default_template(self) -> bytes:
        """
        创建一个简单的默认模板，当其他模板都不可用时使用
        
        Returns:
            bytes: 模板文件的二进制数据
        """
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # 创建文档
            doc = Document()
            
            # 设置页边距
            for section in doc.sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)
            
            # 添加标题
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.add_run("学生综合素质评价报告")
            title_run.font.size = Pt(22)
            title_run.font.bold = True
            
            # 添加学生基本信息
            info_table = doc.add_table(rows=2, cols=4)
            info_table.style = 'Table Grid'
            
            # 填充表格
            cells = [
                ('学生姓名', '{{ 姓名 }}', '学生学号', '{{ 学号 }}'),
                ('班级', '{{ 班级 }}', '学年学期', '{{ 学年 }}{{ 学期 }}')
            ]
            
            for i, row in enumerate(cells):
                for j in range(0, 4, 2):
                    info_table.cell(i, j).text = row[j]
                    info_table.cell(i, j+1).text = row[j+1]
            
            # 添加学生评语部分
            doc.add_paragraph()
            comment_para = doc.add_paragraph()
            comment_para.add_run("学生评语").bold = True
            doc.add_paragraph('{{ 评语 }}')
            
            # 添加学生成绩部分
            doc.add_paragraph()
            grade_para = doc.add_paragraph()
            grade_para.add_run("学习成绩").bold = True
            
            subjects = [
                ('语文', '{{ 语文 }}', '数学', '{{ 数学 }}', '英语', '{{ 英语 }}'),
                ('道法', '{{ 道法 }}', '科学', '{{ 科学 }}', '体育', '{{ 体育 }}'),
                ('音乐', '{{ 音乐 }}', '美术', '{{ 美术 }}', '劳动', '{{ 劳动 }}'),
                ('信息', '{{ 信息 }}', '综合', '{{ 综合 }}', '书法', '{{ 书法 }}')
            ]
            
            grade_table = doc.add_table(rows=len(subjects), cols=6)
            grade_table.style = 'Table Grid'
            
            for i, row in enumerate(subjects):
                for j in range(0, 6, 2):
                    grade_table.cell(i, j).text = row[j]
                    grade_table.cell(i, j+1).text = row[j+1]
            
            # 添加页脚签名
            doc.add_paragraph()
            doc.add_paragraph()
            signature = doc.add_paragraph()
            signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            signature.add_run("教师签名：{{ 教师姓名 }}")
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_para.add_run("日期：{{ 日期 }}")
            
            # 导出为bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"创建默认模板时出错: {str(e)}")
            logger.error(traceback.format_exc())
            # 如果出错，返回一个最小的有效docx文件
            from docx import Document
            doc = Document()
            doc.add_paragraph("学生评价报告")
            doc.add_paragraph("学生姓名: {{ 姓名 }}")
            doc.add_paragraph("评语: {{ 评语 }}")
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
    
    def get_template_path(self, template_id: str) -> str:
        """
        获取模板文件路径
        
        Args:
            template_id: 模板ID，"default"表示默认模板，其他为自定义模板
            
        Returns:
            str: 模板文件的完整路径
        """
        # 特别处理泉州东海湾实验学校模板
        if template_id == "泉州东海湾实验学校综合素质发展报告单":
            return os.path.join(self.templates_dir, "泉州东海湾实验学校综合素质发展报告单.docx")
        
        if template_id == "default":
            return os.path.join(self.templates_dir, "default_template.docx")
        elif template_id and os.path.exists(os.path.join(self.templates_dir, f"{template_id}.docx")):
            return os.path.join(self.templates_dir, f"{template_id}.docx")
        else:
            return os.path.join(self.templates_dir, "custom", f"{template_id}.docx")
    
    def prepare_template_data(self, 
                              student: Dict[str, Any], 
                              comment: Optional[Dict[str, Any]], 
                              grades: Optional[Dict[str, Any]], 
                              settings: Dict[str, Any]) -> Dict[str, Any]:
        """
        准备用于模板替换的数据
        
        Args:
            student: 学生信息
            comment: 学生评语
            grades: 学生成绩
            settings: 导出设置
            
        Returns:
            Dict[str, Any]: 用于模板替换的数据
        """
        # 获取学期文本
        semester_text = "第一学期" if settings.get('semester') == '1' else "第二学期"
        
        # 处理开学时间，转换为"月日"格式
        start_date = settings.get('startDate', '')
        start_date_text = ''
        if start_date:
            try:
                # 尝试解析日期 (格式为YYYY-MM-DD)
                date_parts = start_date.split('-')
                if len(date_parts) == 3:
                    # 移除前导零
                    month = str(int(date_parts[1]))
                    day = str(int(date_parts[2]))
                    start_date_text = f"{month}月{day}日"
                else:
                    start_date_text = start_date
            except Exception as e:
                logger.warning(f"处理开学时间格式出错: {e}")
                start_date_text = start_date
        
        # 安全访问字典的辅助函数
        def safe_get(d, key, default=''):
            if d is None:
                return default
            return d[key] if key in d else default
        
        # 准备用于模板替换的数据 - 注意键名不包含【】，因为已经在模板环境中设置了这些作为分隔符
        data = {
            # 基本学生信息
            "姓名": safe_get(student, 'name'),
            "学号": safe_get(student, 'id'),
            "性别": safe_get(student, 'gender'),
            "班级": safe_get(student, 'class') or settings.get('className', ''),
            
            # 体测数据
            "身高": safe_get(student, 'height'),
            "体重": safe_get(student, 'weight'),
            "肺活量": safe_get(student, 'vital_capacity') or safe_get(student, 'vitalCapacity'),
            "视力左": safe_get(student, 'vision_left') or safe_get(student, 'visionLeft'),
            "视力右": safe_get(student, 'vision_right') or safe_get(student, 'visionRight'),
            "体测情况": safe_get(student, 'physical_test_status') or safe_get(student, 'physicalTestStatus'),
            "胸围": safe_get(student, 'chest_circumference') or safe_get(student, 'chestCircumference'),
            "龋齿": safe_get(student, 'dental_caries') or safe_get(student, 'dentalCaries'),
            
            # 评语
            "评语": safe_get(comment, 'content') if comment else '',
            
            # 报告信息
            "学年": settings.get('schoolYear', ''),
            "学期": semester_text,
            "开学时间": start_date_text,
            "学校名称": settings.get('schoolName', '学校名称未设置'),
            "班主任": settings.get('teacherName', '班主任姓名未设置'),
            "教师姓名": settings.get('teacherName', '教师姓名未设置'),
            "日期": datetime.now().strftime('%Y-%m-%d')
        }
        
        # 如果有成绩数据，添加成绩
        if grades and isinstance(grades, dict) and 'grades' in grades:
            grade_map = {
                'yuwen': '语文',
                'shuxue': '数学',
                'yingyu': '英语',
                'daof': '道法',
                'kexue': '科学',
                'tiyu': '体育',
                'yinyue': '音乐',
                'meishu': '美术',
                'laodong': '劳动',
                'xinxi': '信息',
                'zonghe': '综合',
                'shufa': '书法'
            }
            
            grades_data = grades.get('grades', {})
            for code, placeholder in grade_map.items():
                data[placeholder] = safe_get(grades_data, code)
        
        # 直接从学生数据中获取成绩字段（如果存在于学生数据中）
        grade_fields = ['yuwen', 'shuxue', 'yingyu', 'daof', 'kexue', 'tiyu', 
                      'yinyue', 'meishu', 'laodong', 'xinxi', 'zonghe', 'shufa']
        grade_map = {
            'yuwen': '语文',
            'shuxue': '数学',
            'yingyu': '英语',
            'daof': '道法',
            'kexue': '科学',
            'tiyu': '体育',
            'yinyue': '音乐',
            'meishu': '美术',
            'laodong': '劳动',
            'xinxi': '信息',
            'zonghe': '综合',
            'shufa': '书法'
        }
        
        for field in grade_fields:
            if field in student and field in grade_map and not data.get(grade_map[field]):
                data[grade_map[field]] = safe_get(student, field)
        
        return data
    
    def export_single_report(self, 
                           template_path: str,
                           student: Dict[str, Any],
                           comment: Optional[Dict[str, Any]],
                           grades: Optional[Dict[str, Any]],
                           settings: Dict[str, Any]) -> Tuple[bool, Union[bytes, str]]:
        """
        为单个学生生成报告
        
        Args:
            template_path: 模板文件路径
            student: 学生信息
            comment: 学生评语
            grades: 学生成绩
            settings: 导出设置
            
        Returns:
            Tuple[bool, Union[bytes, str]]: (是否成功, 结果文件内容或错误信息)
        """
        try:
            # 提取学生ID和姓名
            student_id = student.get('id')
            student_name = student.get('name')
            if not student_id or not student_name:
                logger.error(f"学生ID或姓名为空: ID={student_id}, 姓名={student_name}")
                return False, "学生ID或姓名为空"
                
            logger.info(f"开始为学生生成报告: ID={student_id}, 姓名={student_name}")
                
            # 导入docxtpl库，如果不存在则提示安装
            try:
                from docxtpl import DocxTemplate
            except ImportError:
                logger.error("未安装docxtpl库")
                return False, "未安装docxtpl库，请运行 'pip install docxtpl' 安装"
            
            # 检查模板文件是否存在，如果不存在且有内置默认模板，则使用内置模板
            use_builtin_template = False
            if not os.path.exists(template_path):
                logger.warning(f"模板文件不存在: {template_path}")
                if self.has_default_backup:
                    logger.info("将使用内置默认模板")
                    use_builtin_template = True
                else:
                    return False, f"模板文件不存在: {template_path}，且无法创建内置默认模板"
            
            # 准备数据
            context = self.prepare_template_data(student, comment, grades, settings)
            logger.debug(f"准备模板数据完成: {len(context)} 个字段")
            
            # 生成报告
            try:
                if use_builtin_template:
                    # 使用内存中的模板
                    doc_buffer = io.BytesIO(self.default_template)
                    doc = DocxTemplate(doc_buffer)
                else:
                    # 尝试加载模板文件
                    try:
                        doc = DocxTemplate(template_path)
                    except Exception as template_error:
                        logger.error(f"加载模板文件失败: {str(template_error)}")
                        if self.has_default_backup:
                            logger.info("模板加载失败，将使用内置默认模板")
                            doc_buffer = io.BytesIO(self.default_template)
                            doc = DocxTemplate(doc_buffer)
                        else:
                            return False, f"模板文件加载失败: {str(template_error)}"
                
                # 渲染模板
                doc.render(context)
                
                # 保存到内存
                output = io.BytesIO()
                doc.save(output)
                output.seek(0)
                
                return True, output.getvalue()
            except Exception as e:
                error_msg = f"模板渲染错误: {str(e)}"
                logger.error(error_msg)
                logger.error(traceback.format_exc())
                
                # 如果渲染失败但有内置模板，尝试使用内置模板
                if not use_builtin_template and self.has_default_backup:
                    logger.info("模板渲染失败，尝试使用内置默认模板")
                    try:
                        doc_buffer = io.BytesIO(self.default_template)
                        doc = DocxTemplate(doc_buffer)
                        doc.render(context)
                        
                        output = io.BytesIO()
                        doc.save(output)
                        output.seek(0)
                        
                        return True, output.getvalue()
                    except Exception as backup_error:
                        logger.error(f"使用内置默认模板也失败: {str(backup_error)}")
                        return False, f"模板渲染错误: {str(e)}，并且内置默认模板也失败: {str(backup_error)}"
                
                return False, error_msg
                
        except Exception as e:
            error_msg = f"报告生成过程中发生错误: {str(e)}"
            logger.error(error_msg)
            logger.error(traceback.format_exc())
            return False, error_msg
    
    def export_reports(self, 
                       students: List[Dict[str, Any]],
                       comments: Dict[str, Dict[str, Any]],
                       grades: Dict[str, Dict[str, Any]],
                       template_id: str,
                       settings: Dict[str, Any]) -> Tuple[bool, Union[bytes, str]]:
        """
        批量导出多个学生的报告
        
        Args:
            students: 学生信息列表
            comments: 学生评语字典，键为学生ID
            grades: 学生成绩字典，键为学生ID
            template_id: 模板ID
            settings: 导出设置
            
        Returns:
            Tuple[bool, Union[bytes, str]]: (是否成功, 生成的ZIP文件数据或错误消息)
        """
        try:
            # 检查学生数据是否有效
            if not students or not isinstance(students, list) or len(students) == 0:
                logger.error(f"学生数据列表为空或无效")
                return False, "没有提供有效的学生数据"
                
            # 检查模板ID是否有效
            if not template_id:
                logger.error("未提供模板ID")
                if self.has_default_backup:
                    logger.info("未提供模板ID，将使用内置默认模板")
                    template_path = None  # 稍后在处理每个学生时使用内置模板
                else:
                    return False, "未提供模板ID"
            else:
                # 获取模板路径
                template_path = self.get_template_path(template_id)
                # 如果指定模板不存在，但有内置默认模板，使用内置模板
                if not os.path.exists(template_path) and self.has_default_backup:
                    logger.warning(f"模板文件不存在: {template_path}，将使用内置默认模板")
                    template_path = None  # 稍后在处理每个学生时使用内置模板
                elif not os.path.exists(template_path):
                    logger.error(f"模板文件不存在: {template_path}")
                    return False, f"模板文件不存在: {template_path}"
                
            logger.info(f"开始批量导出报告，学生数量: {len(students)}, 使用模板: {template_id}")
            if template_path:
                logger.info(f"模板路径: {template_path}")
            else:
                logger.info("将使用内置默认模板")
                
            # 创建临时目录
            with tempfile.TemporaryDirectory() as temp_dir:
                # 创建ZIP文件
                zip_path = os.path.join(temp_dir, "student_reports.zip")
                
                with zipfile.ZipFile(zip_path, 'w') as zipf:
                    # 遍历学生生成报告
                    success_count = 0
                    error_messages = []
                    
                    for idx, student in enumerate(students):
                        # 安全访问student_id
                        student_id = student.get('id')
                        student_name = student.get('name')
                        
                        if not student_id:
                            logger.warning(f"学生ID为空，跳过该学生: {student}")
                            error_messages.append(f"学生数据不完整: 缺少ID")
                            continue
                        
                        if not student_name:
                            logger.warning(f"学生 {student_id} 缺少姓名，将继续尝试生成报告")
                            
                        logger.info(f"处理第 {idx+1}/{len(students)} 个学生: ID={student_id}, 姓名={student_name or '未知'}")
                            
                        # 获取学生评语和成绩
                        comment = comments.get(student_id, {})
                        grade = grades.get(student_id, {})
                        
                        # 生成报告
                        logger.info(f"为学生 {student_id} 生成报告...")
                        
                        # 传递实际模板路径或None（使用内置模板）
                        success, result = self.export_single_report(
                            template_path or "", student, comment, grade, settings
                        )
                        
                        if success:
                            # 设置文件名格式（默认为 学号_姓名.docx）
                            filename_format = settings.get('fileNameFormat', 'id_name')
                            
                            if filename_format == 'name_id':
                                filename = f"{student_name or '未知'}_{student_id}.docx"
                            elif filename_format == 'id':
                                filename = f"{student_id}.docx"
                            elif filename_format == 'name':
                                filename = f"{student_name or '未知'}.docx"
                            else:  # default: id_name
                                filename = f"{student_id}_{student_name or '未知'}.docx"
                            
                            logger.info(f"报告生成成功，准备保存为: {filename}")
                            
                            # 保存到临时文件
                            file_path = os.path.join(temp_dir, filename)
                            with open(file_path, 'wb') as f:
                                f.write(result)
                            
                            # 添加到ZIP
                            zipf.write(file_path, filename)
                            success_count += 1
                            logger.info(f"成功添加学生 {student_id} 的报告到ZIP文件")
                        else:
                            logger.error(f"生成学生 {student_id} 的报告失败: {result}")
                            error_messages.append(f"学生 {student_id} 报告生成失败: {result}")
                            
                # 读取ZIP文件
                if success_count > 0:
                    logger.info(f"生成ZIP文件: {zip_path}")
                    with open(zip_path, 'rb') as f:
                        zip_data = f.read()
                    
                    logger.info(f"成功导出 {success_count}/{len(students)} 个学生报告")
                    
                    # 如果有部分失败，添加到返回的消息中
                    message = f"成功导出 {success_count}/{len(students)} 个学生报告"
                    if len(error_messages) > 0:
                        message += f"，{len(error_messages)} 个学生报告生成失败"
                    
                    return True, zip_data
                else:
                    error_summary = "\n".join(error_messages[:5])
                    if len(error_messages) > 5:
                        error_summary += f"\n...以及其他 {len(error_messages) - 5} 个错误"
                    
                    logger.error(f"没有成功生成任何报告. 错误: {error_summary}")
                    return False, f"没有成功生成任何报告. 错误: {error_summary}"
                
        except Exception as e:
            logger.error(f"导出报告出错: {str(e)}")
            logger.error(traceback.format_exc())
            return False, f"导出报告失败: {str(e)}"
    
    def save_template(self, template_data: bytes, template_name: str) -> Tuple[bool, str]:
        """
        保存上传的模板文件
        
        Args:
            template_data: 模板文件数据
            template_name: 模板名称
            
        Returns:
            Tuple[bool, str]: (是否成功, 成功消息或错误消息)
        """
        try:
            # 确保模板名称有效
            if not template_name:
                template_name = f"custom_template_{datetime.now().strftime('%Y%m%d%H%M%S')}"
                
            # 确保有.docx后缀
            if not template_name.endswith('.docx'):
                template_name += '.docx'
                
            # 保存文件
            template_path = os.path.join(self.templates_dir, "custom", template_name)
            with open(template_path, 'wb') as f:
                f.write(template_data)
                
            return True, f"模板已保存: {template_name}"
            
        except Exception as e:
            logger.error(f"保存模板出错: {str(e)}")
            return False, f"保存模板失败: {str(e)}"
    
    def list_templates(self) -> List[Dict[str, str]]:
        """
        列出所有可用的模板
        
        Returns:
            List[Dict[str, str]]: 模板列表，每个模板包含id和name
        """
        templates = []
        
        # 添加默认模板
        if os.path.exists(os.path.join(self.templates_dir, "default_template.docx")):
            templates.append({
                "id": "default",
                "name": "默认模板",
                "type": "system"
            })
            
        # 添加自定义模板
        custom_dir = os.path.join(self.templates_dir, "custom")
        if os.path.exists(custom_dir):
            for filename in os.listdir(custom_dir):
                if filename.endswith('.docx'):
                    template_id = os.path.splitext(filename)[0]
                    templates.append({
                        "id": template_id,
                        "name": template_id,
                        "type": "custom"
                    })
                    
        return templates 